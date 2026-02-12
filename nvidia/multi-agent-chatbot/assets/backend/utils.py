#
# SPDX-FileCopyrightText: Copyright (c) 1993-2025 NVIDIA CORPORATION & AFFILIATES. All rights reserved.
# SPDX-License-Identifier: Apache-2.0
#
"""Utility functions for file processing and message conversion."""

import json
import os
import time
from typing import List, Dict, Any, Optional

from langchain_core.documents import Document
from langchain_core.messages import HumanMessage, AIMessage, ToolMessage, ToolCall

from logger import logger
from vector_store import VectorStore


def clean_metadata(metadata: dict) -> dict:
    """Remove problematic metadata fields that Milvus can't handle.
    
    Args:
        metadata: Original metadata dictionary
        
    Returns:
        Cleaned metadata dictionary safe for Milvus
    """
    # Fields that cause "Unrecognized datatype" errors in Milvus
    problematic_fields = {
        'languages', 
        'emphasized_text_contents', 
        'emphasized_text_tags',
        'link_urls',
        'link_texts',
        'parent_id',
        'coordinates',
        'detection_class_prob'
    }
    
    cleaned = {}
    for key, value in metadata.items():
        if key in problematic_fields:
            continue
        # Convert lists/dicts to strings for Milvus compatibility
        if isinstance(value, (list, dict)):
            cleaned[key] = str(value)[:500]  # Limit length
        elif isinstance(value, (int, float, str, bool)) or value is None:
            cleaned[key] = value
        else:
            # Convert any other type to string
            cleaned[key] = str(value)[:500]
    
    return cleaned


async def process_and_ingest_files_background(
    file_info: List[dict],
    vector_store: VectorStore,
    config_manager,
    task_id: str,
    indexing_tasks: Dict[str, str],
    collection: Optional[str] = None,
    doc_type: Optional[str] = None,
) -> None:
    """Process and ingest files in the background.

    Args:
        file_info: List of file dictionaries with 'filename' and 'content' keys
        vector_store: VectorStore instance for document indexing
        config_manager: ConfigManager instance for updating sources
        task_id: Unique identifier for this processing task
        indexing_tasks: Dictionary to track task status
        collection: Optional collection name to ingest into (defaults to vector store's default)
    """
    try:
        logger.debug({
            "message": "Starting background file processing",
            "task_id": task_id,
            "file_count": len(file_info),
            "collection": collection or vector_store.default_collection_name()
        })

        indexing_tasks[task_id] = "saving_files"

        permanent_dir = os.path.join("uploads", task_id)
        os.makedirs(permanent_dir, exist_ok=True)

        file_paths = []
        file_names = []

        for info in file_info:
            try:
                file_name = info["filename"]
                content = info["content"]

                file_path = os.path.join(permanent_dir, file_name)
                with open(file_path, "wb") as f:
                    f.write(content)

                file_paths.append(file_path)
                file_names.append(file_name)

                logger.debug({
                    "message": "Saved file",
                    "task_id": task_id,
                    "filename": file_name,
                    "path": file_path
                })
            except Exception as e:
                logger.error({
                    "message": f"Error saving file {info.get('filename', '<unknown>')}",
                    "task_id": task_id,
                    "filename": info.get('filename', '<unknown>'),
                    "error": str(e)
                }, exc_info=True)

        indexing_tasks[task_id] = "loading_documents"
        logger.debug({"message": "Loading documents", "task_id": task_id})

        try:
            # Use fallback loader with improved error handling
            documents = _fallback_load_documents(file_paths, doc_type=doc_type)

            logger.debug({
                "message": "Documents loaded, starting indexing",
                "task_id": task_id,
                "document_count": len(documents),
                "collection": collection or vector_store.default_collection_name()
            })

            indexing_tasks[task_id] = "indexing_documents"
            vector_store.index_documents(documents, collection_name=collection)

            # Update config sources with the new filenames (if not already present)
            if file_names:
                config = config_manager.read_config()

                config_updated = False
                for file_name in file_names:
                    if file_name not in getattr(config, "sources", []):
                        config.sources.append(file_name)
                        config_updated = True

                if config_updated:
                    config_manager.write_config(config)
                    logger.debug({
                        "message": "Updated config with new sources",
                        "task_id": task_id,
                        "sources": config.sources
                    })

            indexing_tasks[task_id] = "completed"
            logger.debug({
                "message": "Background processing and indexing completed successfully",
                "task_id": task_id
            })
        except Exception as e:
            indexing_tasks[task_id] = f"failed_during_indexing: {str(e)}"
            logger.error({
                "message": "Error during document loading or indexing",
                "task_id": task_id,
                "error": str(e)
            }, exc_info=True)

    except Exception as e:
        indexing_tasks[task_id] = f"failed: {str(e)}"
        logger.error({
            "message": "Error in background processing",
            "task_id": task_id,
            "error": str(e)
        }, exc_info=True)


def _fallback_load_documents(file_paths: List[str], doc_type: Optional[str] = None) -> List[Document]:
    """Minimal, safe loader used only if VectorStore._load_documents is not available."""
    docs: List[Document] = []
    for file_path in file_paths:
        try:
            src_name = os.path.basename(file_path)
            ext = os.path.splitext(file_path)[1].lower()

            text: Optional[str] = None

            # Try Unstructured if available - with cleaned metadata
            try:
                from langchain_unstructured import UnstructuredLoader
                
                loader = UnstructuredLoader(
                    file_path,
                    mode="single",
                    strategy="fast",
                )
                loaded = loader.load()
                for d in loaded:
                    # Normalize and clean metadata
                    md = d.metadata or {}
                    md["source"] = md.get("source") or src_name
                    md["file_path"] = file_path
                    md["filename"] = src_name
                    md = clean_metadata(md)  # ← CRITICAL: Remove problematic fields
                    docs.append(Document(page_content=d.page_content or "", metadata=md))
                logger.debug(f"Successfully loaded {src_name} with UnstructuredLoader")
                continue
            except ImportError:
                logger.debug(f"UnstructuredLoader not available for {src_name}, using fallback")
            except Exception as e:
                logger.warning(f"UnstructuredLoader failed for {src_name}: {str(e)}, using fallback")

            # CSV-specific handling
            if ext == ".csv":
                try:
                    import csv
                    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                        reader = csv.DictReader(f)
                        rows = []
                        for i, row in enumerate(reader):
                            row_text = ", ".join([f"{k}: {v}" for k, v in row.items()])
                            rows.append(f"Row {i+1}: {row_text}")
                        text = "\n".join(rows)
                    logger.debug(f"Successfully loaded CSV {src_name}")
                except Exception as e:
                    logger.warning(f"CSV parsing failed for {src_name}: {str(e)}")
                    text = None

            # DOCX-specific handling
            if ext == ".docx" and text is None:
                try:
                    import docx
                    doc = docx.Document(file_path)
                    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
                    text = "\n\n".join(paragraphs)
                    logger.debug(f"Successfully loaded DOCX {src_name}")
                except Exception as e:
                    logger.warning(f"DOCX parsing failed for {src_name}: {str(e)}")

            # PDF text extraction fallback
            if ext == ".pdf" and text is None:
                try:
                    from pypdf import PdfReader
                    reader = PdfReader(file_path)
                    pages = []
                    for p in reader.pages:
                        try:
                            pages.append(p.extract_text() or "")
                        except Exception:
                            pages.append("")
                    text = "\n\n".join(pages).strip()
                    logger.debug(f"Successfully loaded PDF {src_name}")
                except Exception as e:
                    logger.warning(f"PDF parsing failed for {src_name}: {str(e)}")

            # Plain text read as last resort
            if text is None:
                try:
                    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                        text = f.read()
                    logger.debug(f"Successfully loaded as text {src_name}")
                except Exception as e:
                    logger.warning(f"Text read failed for {src_name}: {str(e)}")
                    text = ""

            # Always produce a Document with clean metadata
            md = {
                "source": src_name,
                "file_path": file_path,
                "filename": src_name,
            }
            if doc_type:
                md["doc_type"] = doc_type
            md = clean_metadata(md)  # Clean even basic metadata
            
            content = (text or "").strip()
            if not content:
                content = f"Document: {src_name}"
            docs.append(Document(page_content=content, metadata=md))
            logger.info(f"Successfully created document for {src_name}")

        except Exception as e:
            logger.error({
                "message": "Fallback loader error",
                "file_path": file_path,
                "error": str(e)
            }, exc_info=True)
            # Create a minimal document to avoid complete failure
            docs.append(Document(
                page_content=f"Error loading document: {os.path.basename(file_path)}",
                metadata=clean_metadata({"source": os.path.basename(file_path), "error": str(e)})
            ))

    return docs


def convert_langgraph_messages_to_openai(messages: List) -> List[Dict[str, Any]]:
    """Convert LangGraph message objects to OpenAI API format.

    Args:
        messages: List of LangGraph message objects

    Returns:
        List of dictionaries in OpenAI API format
    """
    openai_messages = []

    for msg in messages:
        if isinstance(msg, HumanMessage):
            openai_messages.append({
                "role": "user",
                "content": msg.content
            })
        elif isinstance(msg, AIMessage):
            openai_msg = {
                "role": "assistant",
                "content": msg.content or ""
            }
            if hasattr(msg, 'tool_calls') and msg.tool_calls:
                openai_msg["tool_calls"] = []
                for tc in msg.tool_calls:
                    openai_msg["tool_calls"].append({
                        "id": tc["id"],
                        "type": "function",
                        "function": {
                            "name": tc["name"],
                            "arguments": json.dumps(tc["args"])
                        }
                    })
            openai_messages.append(openai_msg)
        elif isinstance(msg, ToolMessage):
            openai_messages.append({
                "role": "tool",
                "content": msg.content,
                "tool_call_id": msg.tool_call_id
            })

    return openai_messages